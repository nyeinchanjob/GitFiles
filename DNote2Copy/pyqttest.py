import sys
import traceback
#import deliverynotes2copy as dn2c

from PyQt4.QtGui import *
from PyQt4.QtCore import *

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xlrd import open_workbook



from lxml import etree
from pprint import pprint

import os
from shutil import copyfile

sourcefilenames = ''
outputfoldername = ''
docxfilename = ''

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(8)

class window(QWidget):
    def __init__(self, parent = None):
        super(window, self).__init__(parent)

        layout = QGridLayout()

        self.labelSource = QLabel("Source : ")
        layout.addWidget(self.labelSource, 0, 0)
        self.textfilepath = QTextEdit()
        self.textfilepath.setReadOnly(True)
        layout.addWidget(self.textfilepath, 0 , 1)
        self.buttonAdd = QPushButton("Add")
        self.buttonAdd.clicked.connect(self.getsourcefile)
        layout.addWidget(self.buttonAdd, 0 , 2)

        self.labelOutput = QLabel("Output : ")
        layout.addWidget(self.labelOutput, 1, 0)
        self.textfolderpath = QTextEdit()
        self.textfolderpath.setReadOnly(True)
        layout.addWidget(self.textfolderpath, 1, 1)
        self.buttonOutput = QPushButton("Add")
        self.buttonOutput.clicked.connect(self.getoutfolder)
        layout.addWidget(self.buttonOutput, 1, 2)

        self.labelFileName = QLabel("File Name : ")
        layout.addWidget(self.labelFileName, 2, 0)
        self.textFileName = QTextEdit()
        layout.addWidget(self.textFileName, 2, 1)

        self.buttonReset = QPushButton("Reset")
        self.buttonReset.clicked.connect(self.reset)
        layout.addWidget(self.buttonReset, 3, 0)
        self.buttonConvert = QPushButton("Convert")
        self.buttonConvert.clicked.connect(self.convert)
        layout.addWidget(self.buttonConvert, 3, 1)
        self.buttonClose = QPushButton("Close")
        self.buttonClose.clicked.connect(self.close)
        layout.addWidget(self.buttonClose, 3, 2)

        self.setLayout(layout)
        self.setWindowTitle("Delivery Notes 2 Copy")
        self.setGeometry(450, 300, 700,100)

    def getsourcefile(self):
        filename = QString()
        filename = QFileDialog.getOpenFileName(self, 'Source file', '', 'Excel 97-2003 file (*.xls)')
        #self.textfilepath.setText(filename.replace('/','\\\\'))
        self.textfilepath.setText(filename)

    def getoutfolder(self):
        outputfoldername = QString()
        outputfoldername = QFileDialog.getExistingDirectory(self, 'Output folder', '', QFileDialog.ShowDirsOnly)
        #self.textfolderpath.setText(outputfoldername + '\\')
        self.textfolderpath.setText(outputfoldername + '/')

    def reset(self):
        self.textfilepath.setText('')
        self.textfolderpath.setText('')

    def convert(self):
        try:
            #orders = dn2c.readxls(self.textfilepath.toPlainText())
            xlsfile = copyFile(self)
            try:

                orders = readxls(xlsfile)
                docxfilename = str(orders[0][2][2]).replace('Delivery Date: ','').replace('.','') + '_' + str(orders[0][4][1]).replace('Sale Rep Name: ','').replace(' ','')
                writedocxwithrealxls(str(self.textfolderpath.toPlainText()), str(self.textFileName.toPlainText()), orders)
            except:
                orders = parseHTML(xlsfile)
                docxfilename = str(orders[0][0][3][1]).replace('Delivery Date: ','').replace('.','') + '_' + str(orders[0][0][3][3]).replace('Sale Rep Name: ','').replace(' ','')
                #dn2c.writedocx(self.textfolderpath.toPlainText(), docxfilename, orders)
                writedocx(str(self.textfolderpath.toPlainText()), str(self.textFileName.toPlainText()), orders)
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Information)
            msg.setWindowTitle("Successful")
            msg.setText("Convert Successfully! Check %s.docx file in selected folder (%s)" % (docxfilename, self.textfolderpath.toPlainText()))
            msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            msg.exec_()
        except:
            print traceback.format_exc()
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Information)
            msg.setWindowTitle("Convert Fail")
            msg.setText("Convert Failed! Check your files or contact your administrator.")
            msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            msg.exec_()



    def checkOutputFolderPath(self):
        createnewfile = True
        file_path = ''
        if os.path.isfile('folderPath.sh'):
            createnewfile = False
            file = open('folderPath.sh', 'r')
            file_path = file.read()
            file.close()
            if len(file_path) == 0:
                file_path = createNewOutputFile(self)
                self.textfolderpath.setText(file_path)
        else:
            #self.textfolderpath.setText(createNewOutputFile() + '\\')
            file_path = createNewOutputFile(self)
            self.textfolderpath.setText(file_path)
        self.textfolderpath.setText(file_path)

    def close(self):
        sys.exit()

def createNewOutputFile(self):
    outputfoldername = QString()
    outputfoldername = QFileDialog.getExistingDirectory(self, 'Output folder', '', QFileDialog.ShowDirsOnly)
    #self.textfolderpath.setText(outputfoldername + '\\')
    file = open('folderPath.sh', 'w')
    file.write(outputfoldername + '/')
    #file.write(outputfoldername + '\\')
    file.close()
    #return outputfoldername + '\\'
    return outputfoldername + '/'

def copyFile(self):
    try:
        copyfile(str(self.textfilepath.toPlainText()), os.path.join(str(self.textfolderpath.toPlainText()), '%s.xls' % (str(self.textFileName.toPlainText()))))
        return os.path.join(str(self.textfolderpath.toPlainText()), '%s.xls' % (str(self.textFileName.toPlainText())))
    except:
        print tracback.format_exc()

def readxls(file_path):
    wb = open_workbook(file_path, on_demand = True)
    values_of_rows = []
    orders = []
    sheet = wb.sheet_by_index(0)
    for row in range(0,sheet.nrows):
        values_of_cells = []
        cells = sheet.row(row)
        for cell in cells:
            if cell.value !='':
                values_of_cells.append(cell.value)
        if len(values_of_cells) > 0:
            values_of_rows.append(values_of_cells)
        if len(values_of_cells) > 0:
            if str(values_of_cells[-1]).replace(' ','').replace('_','') == 'TheSeller':
                orders.append(values_of_rows)
                values_of_rows = []
    orders_detail = []
    total_index = []
    for i in range(0, len(orders)):
        for j in range(0, len(orders[i])):
            if orders[i][j][0] == 'Total Products:':
                orders_details = orders[i][8:j]
                orders[i][8] = orders_details
                total_index.append(len(orders_details) + 1)
    i = 0
    for j in total_index:
        del orders[i][9:7+j]
        i += 1
    return orders

def getFooter(block):
    data = [['Above price already included 5% commercial tax'],
            ['The Buyer  __________________',
             'The Seller  __________________']]
    return data


def getItemDetail(block):
    data = []
    for e, tr in enumerate(block.getchildren()):
        if e == 0:
            for td in tr.getchildren():
                for div in td.getchildren():
                    for table in div.iter("table"):
                        for f, trf in enumerate(table.getchildren()):
                            if f > 0:
                                each_row = []
                                for each_td in trf.iter("td"):
                                    value = each_td.text.strip()
                                    if value:
                                        each_row.append(value)
                                data.append(each_row)
    return data


def getHeader(block):
    data = []
    for each_tr in block.iter("tr"):
        for each_td in each_tr.iter("td"):
            cdata = []
            etd_value = each_td.text.strip()
            if etd_value.lower() not in ["customer name:", "driver message:"]:
                # print etd_value
                if etd_value:
                    data.append(etd_value)
            children = each_td.iter("span")
            for child in children:

                if child.tag not in ["br", "nobr"]:
                    value = (child.text or "").strip()
                    if value:
                        # print '>>', value
                        cdata.append(value)
                    else:
                        # print 'N/A'
                        cdata.append(value)
                        pass

                elif child.tag == "nobr":
                    grand_children = child.getchildren()
                    for each_gc in grand_children:
                        gc_value = each_gc.text.strip()
                        if gc_value:
                            # print '>>>', gc_value
                            cdata.append(gc_value)
            if cdata:
                data.append(cdata)
    return data


def getBlockData(block, index):
    if index == 0:
        return getHeader(block)
    elif index == 2:
        return getItemDetail(block)
    else:
        return getFooter(block)


def parseBlock(block):
    data = []
    if block.tag == "br":
        return None
    else:
        for level1 in block.getchildren():
            if level1.attrib.get('id') != 'printbtn':
                for e, level2 in enumerate(level1.getchildren()):
                    if level2.tag != 'br':
                        #print '>>', e, ' --> ', level2.tag
                        parsed = getBlockData(level2, e)
                        if e == 4:
                            data.extend(parsed)
                        elif e == 2:
                            data.append(parsed[:-1])
                            data.append(parsed[-1])
                        else:
                            data.append(parsed)
    return data


def parseHTML(path):
    data = []
    parser = etree.HTMLParser()
    doc = etree.parse(path, parser)
    root = doc.getroot()
    for body in root.getchildren():
        for each_block in body.getchildren():
            xls_doc = parseBlock(each_block)
            if xls_doc:
                data.append(xls_doc)
    return data



def writedocx(file_path, filename, orders):
    for section in document.sections:
        section.orientation = 1 # 1 is LANDSCAPE, 0 is POTRAIT
        section.page_width = Mm(297) # for A4 Paper
        section.page_height = Mm(210)

        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    for item in orders:
        print item
        table = document.add_table(rows=0, cols=16)
        table.columns[0].width = Inches(0.44)
        table.columns[1].width = Inches(0.75)
        table.columns[2].width = Inches(1.75)
        table.columns[3].width = Inches(0.50)
        table.columns[4].width = Inches(0.55)
        table.columns[5].width = Inches(0.65)
        table.columns[6].width = Inches(0.75)
        table.columns[7].width = Inches(0.1)
        table.columns[8].width = Inches(0.1)
        table.columns[9].width = Inches(0.44)
        table.columns[10].width = Inches(0.75)
        table.columns[11].width = Inches(1.75)
        table.columns[12].width = Inches(0.50)
        table.columns[13].width = Inches(0.55)
        table.columns[14].width = Inches(0.65)
        table.columns[15].width = Inches(0.75)

        #Delivery Notes Title
        row_one = table.add_row().cells
        #row_one[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_one[0].merge(row_one[6])
        row_one[0].paragraphs[0].add_run(item[0][0]).bold = True
        row_one[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Copy
        #row_one[9].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_one[9].merge(row_one[15])
        row_one[9].paragraphs[0].add_run(item[0][0]).bold = True
        row_one[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        #Distributor Name, Customer Name, Order Number
        row_two = table.add_row().cells
        row_two[0].merge(row_two[1])
        row_two[0].text = item[0][1][0]
        row_two[0].paragraphs[0].add_run('\n' + item[0][1][1])
        row_two[2].text = 'Customer Name : '
        row_two[2].paragraphs[0].bold = True
        row_two[2].paragraphs[0].add_run(item[0][2][0])
        row_two[2].paragraphs[0].add_run('\n' + 'Code : ' + item[0][2][1])
        row_two[2].paragraphs[0].add_run('\n' + 'Address : ' + item[0][2][2])
        row_two[2].paragraphs[0].add_run('\n' + 'Tel : ' + item[0][2][3])
        row_two[3].merge(row_two[6])
        row_two[3].text = 'Order Number : '
        row_two[3].paragraphs[0].bold = True
        row_two[3].paragraphs[0].add_run(item[0][3][0])
        row_two[3].paragraphs[0].add_run('\n' + 'Delivery Date : ' + item[0][3][1])
        row_two[3].paragraphs[0].add_run('\n' + 'Geo Code : ' + item[0][3][2])
        row_two[3].paragraphs[0].add_run('\n' + 'Sales Rep Name : ' + item[0][3][3])
        row_two[3].paragraphs[0].add_run('\n' + 'Tel : ' + item[0][3][4])
        #Copy
        row_two[9].merge(row_two[10])
        row_two[9].text = item[0][1][0]
        row_two[9].paragraphs[0].add_run('\n' + item[0][1][1])
        row_two[11].text = 'Customer Name : '
        row_two[11].paragraphs[0].bold = True
        row_two[11].paragraphs[0].add_run(item[0][2][0])
        row_two[11].paragraphs[0].add_run('\n' + 'Code :' + item[0][2][1])
        row_two[11].paragraphs[0].add_run('\n' + 'Address : ' + item[0][2][2])
        row_two[11].paragraphs[0].add_run('\n' + 'Tel : ' + item[0][2][3])
        row_two[12].merge(row_two[15])
        row_two[12].text = 'Order Number : '
        row_two[12].paragraphs[0].bold = True
        row_two[12].paragraphs[0].add_run(item[0][3][0])
        row_two[12].paragraphs[0].add_run('\n' + 'Delivery Date : ' + item[0][3][1])
        row_two[12].paragraphs[0].add_run('\n' + 'Geo Code : ' + item[0][3][2])
        row_two[12].paragraphs[0].add_run('\n' + 'Sales Rep Name : ' + item[0][3][3])
        row_two[0].paragraphs[0].add_run('\n' + 'Tel : ' + item[0][3][4])
        #Driver Message
        row_seven = table.add_row().cells
        row_seven[0].merge(row_seven[6])
        row_seven[0].text = 'DRIVER MESSAGE : ' + item[0][4][0]
        #Copy
        row_seven[9].merge(row_seven[15])
        row_seven[9].text = 'DRIVER MESSAGE : ' + item[0][4][0]
        # Product Detail Title
        row_nine = table.add_row().cells
        row_nine[0].merge(row_nine[6])
        row_table1 = row_nine[0].add_table(rows=0, cols=6)
        row_table1.style = 'TableGrid'
        row_table1.columns[0].width = Inches(0.44)
        row_table1.columns[1].width = Inches(2.0)
        row_table1.columns[2].width = Inches(0.50)
        row_table1.columns[3].width = Inches(0.55)
        row_table1.columns[4].width = Inches(0.65)
        row_table1.columns[5].width = Inches(0.75)
        row_table_cells1 = row_table1.add_row().cells
        row_table_cells1[0].paragraphs[0].add_run('Code').bold = True
        #row_table_cells1[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        #row_table_cells1[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[1].paragraphs[0].add_run('Description').bold = True
        row_table_cells1[2].paragraphs[0].add_run('UOM').bold = True
        row_table_cells1[3].paragraphs[0].add_run('QTY').bold = True
        row_table_cells1[4].paragraphs[0].add_run('Price').bold = True
        row_table_cells1[5].paragraphs[0].add_run('Amount').bold = True
        row_table_cells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        #Copy
        row_nine[9].merge(row_nine[15])
        row_table2 = row_nine[9].add_table(rows=0, cols=6)
        row_table2.style = 'TableGrid'
        row_table2.columns[0].width = Inches(0.44)
        row_table2.columns[1].width = Inches(2.0)
        row_table2.columns[2].width = Inches(0.50)
        row_table2.columns[3].width = Inches(0.55)
        row_table2.columns[4].width = Inches(0.65)
        row_table2.columns[5].width = Inches(0.75)
        row_table_cells2 = row_table2.add_row().cells
        row_table_cells2[0].paragraphs[0].add_run('Code').bold = True
        #row_table_cells2[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        #row_table_cells2[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[1].paragraphs[0].add_run('Description').bold = True
        row_table_cells2[2].paragraphs[0].add_run('UOM').bold = True
        row_table_cells2[3].paragraphs[0].add_run('QTY').bold = True
        row_table_cells2[4].paragraphs[0].add_run('Price').bold = True
        row_table_cells2[5].paragraphs[0].add_run('Amount').bold = True
        row_table_cells2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for products in item[1]:
            #Product Detail
            row_product1 = row_table1.add_row().cells
            row_product1[0].text = str(products[0]).replace('.0','')
            row_product1[1].text = str(products[1])
            row_product1[2].text = str(products[2])
            row_product1[3].text = str(products[3]).replace('.0','')
            row_product1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_product1[4].text = str(products[4]).replace('.0','')
            row_product1[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_product1[5].text = str(products[5]).replace('.0','')
            row_product1[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #Copy
            row_product2 = row_table2.add_row().cells
            row_product2[0].text = str(products[0]).replace('.0','')
            row_product2[1].text = str(products[1])
            row_product2[2].text = str(products[2])
            row_product2[3].text = str(products[3]).replace('.0','')
            row_product2[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_product2[4].text = str(products[4]).replace('.0','')
            row_product2[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_product2[5].text = str(products[5]).replace('.0','')
            row_product2[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        #Total Products
        row_eleven1 = row_table1.add_row().cells
        row_eleven1[0].text = ''
        row_eleven1[1].text = item[2][0]
        row_eleven1[2].text = ''
        row_eleven1[3].text = str(item[2][1]).replace('.0','')
        row_eleven1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_eleven1[4].text = ''
        row_eleven1[5].text = str(item[2][2]).replace('.0','')
        row_eleven1[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #Copy
        row_eleven2 = row_table2.add_row().cells
        row_eleven2[0].text = ''
        row_eleven2[1].text = item[2][0]
        row_eleven2[2].text = ''
        row_eleven2[3].text = str(item[2][1]).replace('.0','')
        row_eleven2[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_eleven2[4].text = ''
        row_eleven2[5].text = str(item[2][2]).replace('.0','')
        row_eleven2[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        #Tax Informaiton
        row_twelve = table.add_row().cells
        row_twelve[0].merge(row_twelve[6])
        row_twelve[0].text = item[3]
        #Copy
        row_twelve[9].merge(row_twelve[15])
        row_twelve[9].text = item[3]

        row_thirteen = table.add_row().cells
        row_fourteen = table.add_row().cells
        row_fifteen = table.add_row().cells
        #Buyer, Seller
        row_sixteen = table.add_row().cells
        row_sixteen[0].merge(row_sixteen[2])
        row_sixteen[0].text = item[4][0]
        row_sixteen[4].merge(row_sixteen[6])
        row_sixteen[4].text = item[4][1]
        #Copy
        row_sixteen[9].merge(row_sixteen[11])
        row_sixteen[9].text = item[4][0]
        row_sixteen[13].merge(row_sixteen[15])
        row_sixteen[13].text = item[4][1]
        document.add_page_break()
    document.save('%s%s.docx' % (file_path, filename))


def writedocxwithrealxls(file_path, filename, orders):

    for section in document.sections:
        section.orientation = 1 # 1 is LANDSCAPE, 0 is POTRAIT
        section.page_width = Mm(297) # for A4 Paper
        section.page_height = Mm(210)

        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)


    for item in orders:
        table = document.add_table(rows=0, cols=16)
        table.columns[0].width = Inches(0.44)
        table.columns[1].width = Inches(0.75)
        table.columns[2].width = Inches(1.75)
        table.columns[3].width = Inches(0.50)
        table.columns[4].width = Inches(0.55)
        table.columns[5].width = Inches(0.65)
        table.columns[6].width = Inches(0.75)
        table.columns[7].width = Inches(0.1)
        table.columns[8].width = Inches(0.1)
        table.columns[9].width = Inches(0.44)
        table.columns[10].width = Inches(0.75)
        table.columns[11].width = Inches(1.75)
        table.columns[12].width = Inches(0.50)
        table.columns[13].width = Inches(0.55)
        table.columns[14].width = Inches(0.65)
        table.columns[15].width = Inches(0.75)


        #Delivery Notes Title
        row_one = table.add_row().cells
        row_one[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_one[0].merge(row_one[6])
        row_one[0].paragraphs[0].add_run(item[0][0]).bold = True

        # Copy
        row_one[9].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_one[9].merge(row_one[15])
        row_one[9].paragraphs[0].add_run(item[0][0]).bold = True

        #Distributor Name, Customer Name, Order Number
        row_two = table.add_row().cells
        row_two[0].merge(row_two[1])
        row_two[0].text = item[1][0]
        row_two[0].paragraphs[0].add_run('\n' + item[2][0])
        row_two[2].text = item[1][1]
        row_two[2].paragraphs[0].add_run('\n' + item[2][1])
        row_two[2].paragraphs[0].add_run('\n' + item[3][0])
        row_two[2].paragraphs[0].add_run('\n' + item[4][0])
        row_two[3].merge(row_two[6])
        row_two[3].text = item[1][2]
        row_two[3].paragraphs[0].add_run('\n' + item[2][2])
        row_two[3].paragraphs[0].add_run('\n' + item[3][1])
        row_two[3].paragraphs[0].add_run('\n' + item[4][1])
        row_two[3].paragraphs[0].add_run('\n' + item[5][0])
        #Copy
        row_two[9].merge(row_two[10])
        row_two[9].text = item[1][0]
        row_two[9].paragraphs[0].add_run('\n' + item[2][0])
        row_two[11].text = item[1][1]
        row_two[11].paragraphs[0].add_run('\n' + item[2][1])
        row_two[11].paragraphs[0].add_run('\n' + item[3][0])
        row_two[11].paragraphs[0].add_run('\n' + item[4][0])
        row_two[12].merge(row_two[15])
        row_two[12].text = item[1][2]
        row_two[12].paragraphs[0].add_run('\n' + item[2][2])
        row_two[12].paragraphs[0].add_run('\n' + item[3][1])
        row_two[12].paragraphs[0].add_run('\n' + item[4][1])
        row_two[0].paragraphs[0].add_run('\n' + item[5][0])
        #Driver Message
        row_seven = table.add_row().cells
        row_seven[0].merge(row_seven[6])
        row_seven[0].text = item[6][0]
        #Copy
        row_seven[9].merge(row_seven[15])
        row_seven[9].text = item[6][0]
        # Product Detail Title
        row_nine = table.add_row().cells
        row_nine[0].merge(row_nine[6])
        row_table1 = row_nine[0].add_table(rows=0, cols=6)
        row_table1.style = 'TableGrid'
        row_table1.columns[0].width = Inches(0.44)
        row_table1.columns[1].width = Inches(2.0)
        row_table1.columns[2].width = Inches(0.50)
        row_table1.columns[3].width = Inches(0.55)
        row_table1.columns[4].width = Inches(0.65)
        row_table1.columns[5].width = Inches(0.75)
        row_table_cells1 = row_table1.add_row().cells
        row_table_cells1[0].paragraphs[0].add_run('Code').bold = True
        row_table_cells1[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells1[1].paragraphs[0].add_run('Description').bold = True
        row_table_cells1[2].paragraphs[0].add_run('UOM').bold = True
        row_table_cells1[3].paragraphs[0].add_run('QTY').bold = True
        row_table_cells1[4].paragraphs[0].add_run('Price').bold = True
        row_table_cells1[5].paragraphs[0].add_run('Amount').bold = True
        #Copy
        row_nine[9].merge(row_nine[15])
        row_table2 = row_nine[9].add_table(rows=0, cols=6)
        row_table2.style = 'TableGrid'
        row_table2.columns[0].width = Inches(0.44)
        row_table2.columns[1].width = Inches(2.0)
        row_table2.columns[2].width = Inches(0.50)
        row_table2.columns[3].width = Inches(0.55)
        row_table2.columns[4].width = Inches(0.65)
        row_table2.columns[5].width = Inches(0.75)
        row_table_cells2 = row_table2.add_row().cells
        row_table_cells2[0].paragraphs[0].add_run('Code').bold = True
        row_table_cells2[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_table_cells2[1].paragraphs[0].add_run('Description').bold = True
        row_table_cells2[2].paragraphs[0].add_run('UOM').bold = True
        row_table_cells2[3].paragraphs[0].add_run('QTY').bold = True
        row_table_cells2[4].paragraphs[0].add_run('Price').bold = True
        row_table_cells2[5].paragraphs[0].add_run('Amount').bold = True

        for products in item[8]:
            #Product Detail
            row_product1 = row_table1.add_row().cells
            row_product1[0].text = str(products[0]).replace('.0','')
            row_product1[1].text = str(products[1])
            row_product1[2].text = str(products[2])
            row_product1[3].text = str(products[3]).replace('.0','')
            row_product1[4].text = str(products[4]).replace('.0','')
            row_product1[5].text = str(products[5]).replace('.0','')
            #Copy
            row_product2 = row_table2.add_row().cells
            row_product2[0].text = str(products[0]).replace('.0','')
            row_product2[1].text = str(products[1])
            row_product2[2].text = str(products[2])
            row_product2[3].text = str(products[3]).replace('.0','')
            row_product2[4].text = str(products[4]).replace('.0','')
            row_product2[5].text = str(products[5]).replace('.0','')

        #Total Products
        row_eleven1 = row_table1.add_row().cells
        row_eleven1[0].text = ''
        row_eleven1[1].text = item[9][0]
        row_eleven1[2].text = ''
        row_eleven1[3].text = str(item[9][1]).replace('.0','')
        row_eleven1[4].text = ''
        row_eleven1[5].text = str(item[9][2]).replace('.0','')
        #Copy
        row_eleven2 = row_table2.add_row().cells
        row_eleven2[0].text = ''
        row_eleven2[1].text = item[9][0]
        row_eleven2[2].text = ''
        row_eleven2[3].text = str(item[9][1]).replace('.0','')
        row_eleven2[4].text = ''
        row_eleven2[5].text = str(item[9][2]).replace('.0','')

        #Tax Informaiton
        row_twelve = table.add_row().cells
        row_twelve[0].merge(row_twelve[6])
        row_twelve[0].text = item[10][0]
        #Copy
        row_twelve[9].merge(row_twelve[15])
        row_twelve[9].text = item[10][0]

        row_thirteen = table.add_row().cells
        row_fourteen = table.add_row().cells
        row_fifteen = table.add_row().cells
        #Buyer, Seller
        row_sixteen = table.add_row().cells
        row_sixteen[0].merge(row_sixteen[2])
        row_sixteen[0].text = item[11][0]
        row_sixteen[4].merge(row_sixteen[6])
        row_sixteen[4].text = item[11][1]
        #Copy
        row_sixteen[9].merge(row_sixteen[11])
        row_sixteen[9].text = item[11][0]
        row_sixteen[13].merge(row_sixteen[15])
        row_sixteen[13].text = item[11][1]

        document.add_page_break()

    document.save('%s%s.docx' % (file_path, filename))

def main():
    app = QApplication(sys.argv)
    ex = window()
    ex.checkOutputFolderPath()
    ex.show()

    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
